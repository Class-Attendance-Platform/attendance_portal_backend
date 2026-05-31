"""
Export generators for attendance reports.
Each function receives a CourseInfo object + optional date filter,
and returns a Django HttpResponse with the correct Content-Type.
"""
import csv
import io
from datetime import date

from django.http import HttpResponse
from django.utils import timezone


def _get_report_data(course_info, filter_date=None):
    """
    Shared data-fetching logic for all exporters.
    Returns: (course_info, students, logs_by_student, all_dates)
    """
    from apps.academic.models import StudentClassroom
    from apps.attendance.models import AttendanceLog

    # All students in this course's classroom
    memberships = StudentClassroom.objects.filter(
        classroom=course_info.classroom
    ).select_related('student__user').order_by('student__student_id')
    students = [m.student for m in memberships]

    # Fetch logs
    logs_qs = AttendanceLog.objects.filter(course_info=course_info).order_by('date')
    if filter_date:
        logs_qs = logs_qs.filter(date=filter_date)

    # Build lookup: { student_id: { date_str: status } }
    logs_by_student = {}
    all_dates = sorted({str(log.date) for log in logs_qs})
    for log in logs_qs:
        sid = log.student.student_id
        if sid not in logs_by_student:
            logs_by_student[sid] = {}
        logs_by_student[sid][str(log.date)] = log.status

    return students, logs_by_student, all_dates


# ── CSV ───────────────────────────────────────────────────────────────────────

def export_csv(course_info, filter_date=None):
    students, logs_by_student, all_dates = _get_report_data(course_info, filter_date)

    response = HttpResponse(content_type='text/csv')
    filename = _filename(course_info, filter_date, 'csv')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    writer = csv.writer(response)

    # Header row
    header = ['Student ID', 'Name', 'Email']
    if filter_date:
        header += ['Status']
    else:
        header += all_dates + ['Total Classes', 'Present', 'Percentage']
    writer.writerow(header)

    for student in students:
        sid = student.student_id
        row = [sid, student.user.get_full_name(), student.user.email]
        student_logs = logs_by_student.get(sid, {})

        if filter_date:
            row.append(student_logs.get(str(filter_date), 'ABSENT'))
        else:
            for d in all_dates:
                row.append(student_logs.get(d, 'ABSENT'))
            total   = len(all_dates)
            present = sum(1 for d in all_dates if student_logs.get(d) == 'PRESENT')
            pct     = round(present / total * 100, 2) if total else 0.0
            row += [total, present, f'{pct}%']

        writer.writerow(row)

    return response


# ── XLSX ──────────────────────────────────────────────────────────────────────

def export_xlsx(course_info, filter_date=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    students, logs_by_student, all_dates = _get_report_data(course_info, filter_date)

    wb = Workbook()
    ws = wb.active
    ws.title = 'Attendance Report'

    # Styles
    header_font    = Font(bold=True, color='FFFFFF', size=11)
    header_fill    = PatternFill(fill_type='solid', fgColor='1F4E79')
    present_fill   = PatternFill(fill_type='solid', fgColor='C6EFCE')
    absent_fill    = PatternFill(fill_type='solid', fgColor='FFC7CE')
    late_fill      = PatternFill(fill_type='solid', fgColor='FFEB9C')
    center_align   = Alignment(horizontal='center', vertical='center')
    thin_border    = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    # Title row
    title = f"{course_info.course.code} — {course_info.course.title}"
    if filter_date:
        title += f"  |  {filter_date}"
    ws.merge_cells(f'A1:{get_column_letter(4 + len(all_dates))}1')
    title_cell = ws['A1']
    title_cell.value = title
    title_cell.font  = Font(bold=True, size=13, color='1F4E79')
    title_cell.alignment = center_align
    ws.row_dimensions[1].height = 22

    # Header row
    headers = ['#', 'Student ID', 'Name', 'Email']
    if filter_date:
        headers += ['Status']
    else:
        headers += all_dates + ['Total', 'Present', '%']

    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col_idx, value=h)
        cell.font      = header_font
        cell.fill      = header_fill
        cell.alignment = center_align
        cell.border    = thin_border

    # Data rows
    for row_idx, student in enumerate(students, 1):
        sid         = student.student_id
        student_logs = logs_by_student.get(sid, {})
        excel_row   = row_idx + 2

        base_values = [row_idx, sid, student.user.get_full_name(), student.user.email]
        for col_idx, val in enumerate(base_values, 1):
            cell = ws.cell(row=excel_row, column=col_idx, value=val)
            cell.border    = thin_border
            cell.alignment = center_align if col_idx != 3 else Alignment(horizontal='left', vertical='center')

        if filter_date:
            status = student_logs.get(str(filter_date), 'ABSENT')
            cell = ws.cell(row=excel_row, column=5, value=status)
            cell.border    = thin_border
            cell.alignment = center_align
            cell.fill = present_fill if status == 'PRESENT' else (late_fill if status == 'LATE' else absent_fill)
        else:
            total   = len(all_dates)
            present_count = 0
            for d_idx, d in enumerate(all_dates):
                status = student_logs.get(d, 'ABSENT')
                if status == 'PRESENT':
                    present_count += 1
                cell = ws.cell(row=excel_row, column=5 + d_idx, value=status)
                cell.border    = thin_border
                cell.alignment = center_align
                cell.fill = present_fill if status == 'PRESENT' else (late_fill if status == 'LATE' else absent_fill)

            pct = round(present_count / total * 100, 2) if total else 0.0
            offset = 5 + len(all_dates)
            for col_idx, val in enumerate([total, present_count, f'{pct}%'], offset):
                cell = ws.cell(row=excel_row, column=col_idx, value=val)
                cell.border    = thin_border
                cell.alignment = center_align

    # Auto column widths
    for col in ws.columns:
        max_len = max((len(str(c.value or '')) for c in col), default=8)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 30)

    # Stream response
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{_filename(course_info, filter_date, "xlsx")}"'
    return response


# ── PDF ───────────────────────────────────────────────────────────────────────

def export_pdf(course_info, filter_date=None):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    students, logs_by_student, all_dates = _get_report_data(course_info, filter_date)

    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm,
    )

    styles  = getSampleStyleSheet()
    story   = []

    # Title
    title_style = ParagraphStyle('title', parent=styles['Title'], fontSize=14, spaceAfter=6)
    subtitle    = f"{course_info.course.code}  —  {course_info.course.title}"
    if filter_date:
        subtitle += f"  |  {filter_date}"
    story.append(Paragraph(subtitle, title_style))
    story.append(Paragraph(
        f"Teacher: {course_info.teacher.user.get_full_name() if course_info.teacher else 'N/A'}  "
        f"|  Classroom: {course_info.classroom.name}  "
        f"|  Generated: {timezone.now().strftime('%Y-%m-%d %H:%M')}",
        styles['Normal']
    ))
    story.append(Spacer(1, 0.4*cm))

    # Table headers
    if filter_date:
        headers = ['#', 'Student ID', 'Name', 'Status']
    else:
        headers = ['#', 'Student ID', 'Name'] + all_dates + ['Total', 'Present', '%']

    table_data = [headers]

    PRESENT_COLOR = colors.HexColor('#C6EFCE')
    ABSENT_COLOR  = colors.HexColor('#FFC7CE')
    LATE_COLOR    = colors.HexColor('#FFEB9C')

    cell_styles = []

    for row_idx, student in enumerate(students, 1):
        sid          = student.student_id
        student_logs = logs_by_student.get(sid, {})
        row          = [row_idx, sid, student.user.get_full_name()]

        if filter_date:
            status = student_logs.get(str(filter_date), 'ABSENT')
            row.append(status)
            col = 3  # 0-indexed
            tbl_row = row_idx + 1
            fill = PRESENT_COLOR if status == 'PRESENT' else (LATE_COLOR if status == 'LATE' else ABSENT_COLOR)
            cell_styles.append(('BACKGROUND', (col, tbl_row), (col, tbl_row), fill))
        else:
            total = len(all_dates)
            present_count = 0
            for d_idx, d in enumerate(all_dates):
                status = student_logs.get(d, 'ABSENT')
                if status == 'PRESENT':
                    present_count += 1
                row.append(status)
                col     = 3 + d_idx
                tbl_row = row_idx + 1
                fill = PRESENT_COLOR if status == 'PRESENT' else (LATE_COLOR if status == 'LATE' else ABSENT_COLOR)
                cell_styles.append(('BACKGROUND', (col, tbl_row), (col, tbl_row), fill))

            pct = round(present_count / total * 100, 2) if total else 0.0
            row += [total, present_count, f'{pct}%']

        table_data.append(row)

    # Build table
    tbl = Table(table_data, repeatRows=1)
    base_style = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E79')),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0, 0), (-1, 0), 9),
        ('FONTSIZE',   (0, 1), (-1, -1), 8),
        ('GRID',       (0, 0), (-1, -1), 0.4, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')]),
        ('ALIGN',      (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN',      (2, 1), (2, -1), 'LEFT'),
        ('VALIGN',     (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ] + cell_styles

    tbl.setStyle(TableStyle(base_style))
    story.append(tbl)
    doc.build(story)

    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{_filename(course_info, filter_date, "pdf")}"'
    return response


# ── DOCX ──────────────────────────────────────────────────────────────────────

def export_docx(course_info, filter_date=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    students, logs_by_student, all_dates = _get_report_data(course_info, filter_date)

    doc = Document()

    # Page layout — landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin  = section.bottom_margin = Cm(1.5)

    # Title
    title_para = doc.add_heading('', level=1)
    run = title_para.add_run(f"{course_info.course.code} — {course_info.course.title}")
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_line = (
        f"Teacher: {course_info.teacher.user.get_full_name() if course_info.teacher else 'N/A'} "
        f"| Classroom: {course_info.classroom.name} "
        f"| Generated: {timezone.now().strftime('%Y-%m-%d %H:%M')}"
    )
    if filter_date:
        info_line += f" | Date: {filter_date}"
    info_para = doc.add_paragraph(info_line)
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Table
    if filter_date:
        col_headers = ['#', 'Student ID', 'Name', 'Status']
    else:
        col_headers = ['#', 'Student ID', 'Name'] + all_dates + ['Total', 'Present', '%']

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr_cells[i].text = str(h)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr_cells[i], '1F4E79')
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    COLORS = {'PRESENT': 'C6EFCE', 'ABSENT': 'FFC7CE', 'LATE': 'FFEB9C'}

    for row_idx, student in enumerate(students, 1):
        sid          = student.student_id
        student_logs = logs_by_student.get(sid, {})
        row_cells    = table.add_row().cells

        base = [str(row_idx), str(sid), student.user.get_full_name()]
        for i, val in enumerate(base):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER
            )

        if filter_date:
            status = student_logs.get(str(filter_date), 'ABSENT')
            row_cells[3].text = status
            _set_cell_bg(row_cells[3], COLORS.get(status, 'FFFFFF'))
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            total = len(all_dates)
            present_count = 0
            for d_idx, d in enumerate(all_dates):
                status = student_logs.get(d, 'ABSENT')
                if status == 'PRESENT':
                    present_count += 1
                col_cell = row_cells[3 + d_idx]
                col_cell.text = status
                _set_cell_bg(col_cell, COLORS.get(status, 'FFFFFF'))
                col_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            pct = round(present_count / total * 100, 2) if total else 0.0
            offset = 3 + len(all_dates)
            for i, val in enumerate([str(total), str(present_count), f'{pct}%'], offset):
                row_cells[i].text = val
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{_filename(course_info, filter_date, "docx")}"'
    return response


# ── Helpers ───────────────────────────────────────────────────────────────────

def _filename(course_info, filter_date, ext):
    code = course_info.course.code.replace(' ', '_')
    if filter_date:
        return f'{code}_attendance_{filter_date}.{ext}'
    return f'{code}_attendance_full.{ext}'


def _set_cell_bg(cell, hex_color: str):
    """Set background color of a docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)
